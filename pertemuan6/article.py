from bs4 import BeautifulSoup
import fungsi1
import requests
import os
from docx import Document  # tambahan untuk file Word

def main_scraper(url, directory):
    fungsi1.create_directory(directory)  # Membuat Directory
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
    }

    source_code = requests.get(url, headers=headers)
    source_text = source_code.text
    soup = BeautifulSoup(source_text, 'html.parser')
    
    # Ambil semua artikel di halaman Kompas Tekno Gadget
    articles = soup.find_all("div", class_="article__box")

    # File teks dan file Word
    file_path = os.path.join(directory, "hasil_gadget.txt")
    file_doc = os.path.join(directory, "hasil_gadget.doc")

    # Buat dokumen Word baru
    doc = Document()
    doc.add_heading("Hasil Scraping Kompas Tekno - Gadget", level=1)

    with open(file_path, "w", encoding="utf-8") as file:
        for article in articles:
            # Ambil judul
            title_tag = article.find("h3", class_="article__title")
            if not title_tag:
                continue
            judul = title_tag.get_text(strip=True)
            
            # Ambil link
            url_gambar = title_tag.a["href"] if title_tag.a else "Tidak ada link"

            # Ambil tanggal
            tanggal_tag = article.find("div", class_="article__date")
            tanggal = tanggal_tag.get_text(strip=True) if tanggal_tag else "Tidak ada tanggal"

            # Tulis ke file txt
            file.write(f"Judul  : {judul}\n")
            file.write(f"URL    : {url_gambar}\n")
            file.write(f"Tanggal: {tanggal}\n\n")

            # Cetak ke terminal
            print("Judul  :", judul)
            print("URL    :", url_gambar)
            print("Tanggal:", tanggal)
            print()

            # Tambah ke file Word
            doc.add_heading(judul, level=2)
            doc.add_paragraph(f"URL    : {url_gambar}")
            doc.add_paragraph(f"Tanggal: {tanggal}")
            doc.add_paragraph("")  # spasi antar artikel

    # Simpan file Word
    doc.save(file_doc)
    print(f"\n Hasil scraping disimpan ke:")
    print(f" {file_path}")
    print(f" {file_doc}")

    # Setelah menulis, baca isi file txt
    read_data(file_path)

def read_data(path):
    with open(path, 'rt', encoding='utf-8') as file:
        for line in file:
            print(line.replace("\n", ""))

# Jalankan
main_scraper("https://tekno.kompas.com/gadget", "Hasil")
